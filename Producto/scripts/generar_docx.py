# -*- coding: utf-8 -*-
"""
Script para generar el documento DOCX de Solución Digital
Proyecto: Césped Sintético Chile
Ejecutar: python generar_docx.py
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime
except ImportError:
    print("Instalando python-docx...")
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

# ── Colores corporativos ──────────────────────────────────────
COLOR_VERDE     = RGBColor(0x2C, 0x59, 0x26)   # #2c5926
COLOR_VERDE_CLARO = RGBColor(0xD6, 0xEA, 0xF8)
COLOR_BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GRIS      = RGBColor(0x64, 0x74, 0x8B)
COLOR_NEGRO     = RGBColor(0x1E, 0x29, 0x3B)

NOMBRE_ARCHIVO = "Solucion_Digital_CespedSintetico.docx"

# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Pinta el fondo de una celda con el color hex dado."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def agregar_heading(doc, texto, nivel=1, color=None):
    h = doc.add_heading(texto, level=nivel)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color or COLOR_VERDE
        run.font.bold = True
    return h

def agregar_parrafo(doc, texto, negrita=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.font.bold = negrita
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p

def agregar_bullet(doc, texto, nivel=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(texto)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5 + nivel * 0.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def tabla_dos_col(doc, filas_data, header=None, col_widths=(5.5, 8.5)):
    """Crea una tabla de 2 columnas con datos."""
    cols = 2
    table = doc.add_table(rows=0, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if header:
        row = table.add_row()
        for i, h in enumerate(header):
            cell = row.cells[i]
            cell.width = Cm(col_widths[i])
            set_cell_bg(cell, '2C5926')
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.bold = True
            run.font.color.rgb = COLOR_BLANCO
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for fila in filas_data:
        row = table.add_row()
        for i, val in enumerate(fila):
            cell = row.cells[i]
            cell.width = Cm(col_widths[i])
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = COLOR_VERDE
    return table

def linea_separadora(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2C5926')
    pb.append(bottom)
    pPr.append(pb)
    return p

# ─────────────────────────────────────────────────────────────
# DOCUMENTO PRINCIPAL
# ─────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# Fuente por defecto
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ══════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = titulo.add_run("SOLUCIÓN DIGITAL")
t_run.font.size = Pt(32)
t_run.font.bold = True
t_run.font.color.rgb = COLOR_VERDE

subtitulo = doc.add_paragraph()
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitulo.add_run("Plataforma Web Césped Sintético Chile")
sub_run.font.size = Pt(18)
sub_run.font.color.rgb = COLOR_GRIS

doc.add_paragraph()
linea_separadora(doc)
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
fecha_actual = datetime.date.today().strftime("%B %Y").capitalize()
meta_run = meta.add_run(f"Versión 1.0  ·  {fecha_actual}  ·  Estado: Desarrollo Activo")
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = COLOR_GRIS
meta_run.font.italic = True

doc.add_page_break()

# ══════════════════════════════════════════
# SECCIÓN 1 – VISIÓN GENERAL
# ══════════════════════════════════════════
agregar_heading(doc, "1. Visión General", 1)
linea_separadora(doc)
doc.add_paragraph()

agregar_parrafo(doc,
    "Césped Sintético Chile es una plataforma web comercial e inteligente diseñada para digitalizar "
    "completamente el ciclo de venta y atención al cliente de una empresa especializada en instalación "
    "y comercialización de pasto sintético para ambientes deportivos, residenciales y de jardinería.")

agregar_parrafo(doc,
    "La plataforma trasciende el concepto de sitio web estático para convertirse en un ecosistema digital "
    "de 5 módulos integrados que trabajan en conjunto para captar clientes, orientarlos en su proceso de "
    "compra, responder sus dudas de forma autónoma y gestionar la trazabilidad de documentos comerciales "
    "(cotizaciones y ventas).")

doc.add_paragraph()

# ══════════════════════════════════════════
# SECCIÓN 2 – PROPUESTA DE VALOR
# ══════════════════════════════════════════
agregar_heading(doc, "2. Propuesta de Valor", 1)
linea_separadora(doc)
doc.add_paragraph()

tabla_dos_col(doc,
    filas_data=[
        ("Clientes sin orientación sobre tipo de pasto",          "Catálogo segmentado: Deportivo, Jardín, Insumos"),
        ("Tiempo perdido cotizando proyectos simples",             "Calculadora interactiva de metraje y presupuesto"),
        ("Consultas frecuentes que ocupan asesores",              "Chatbot con IA conectado a N8N — responde 24/7"),
        ("Clientes que no recuerdan número de cotización",        "Buscador de documentos por RUT o código de folio"),
        ("RUT inválidos generan errores en el sistema back-end",  "Validador en tiempo real con algoritmo Módulo 11 (SII)"),
    ],
    header=["Problema del Negocio", "Solución Digital Implementada"],
    col_widths=(7, 8)
)
doc.add_paragraph()

# ══════════════════════════════════════════
# SECCIÓN 3 – MÓDULOS DEL SISTEMA
# ══════════════════════════════════════════
agregar_heading(doc, "3. Módulos del Sistema", 1)
linea_separadora(doc)
doc.add_paragraph()

# Módulo 1
agregar_heading(doc, "3.1 Módulo 1 — Vitrina Digital y Catálogo", 2)
agregar_parrafo(doc, "La vitrina es el punto de entrada principal del cliente. Compuesta por páginas HTML dedicadas a cada línea de negocio:")
for item in [
    "index.html — Página de inicio con propuesta de valor y llamados a la acción.",
    "pasto-deportivo.html — Catálogo deportivo: canchas de fútbol, pádel y tenis.",
    "jardines.html — Catálogo residencial y decorativo.",
    "insumos.html — Materiales de instalación (adhesivos, caucho granulado, arena de sílice, cinta de unión).",
    "contacto.html — Formulario de contacto directo con el equipo comercial.",
]:
    agregar_bullet(doc, item)

doc.add_paragraph()

# Módulo 2
agregar_heading(doc, "3.2 Módulo 2 — Herramientas de Autogestión (Calculadora)", 2)
agregar_parrafo(doc,
    "Permite al visitante estimar costos de su proyecto ingresando metros cuadrados, tipo de pasto y "
    "opciones de instalación. Reduce la fricción en el proceso de venta al entregar prospectos informados.")

doc.add_paragraph()

# Módulo 3
agregar_heading(doc, "3.3 Módulo 3 — Asistente Conversacional Inteligente (Chatbot)", 2)
agregar_parrafo(doc,
    'El chatbot "Queno" opera en todas las páginas del sitio con persistencia de historial vía sessionStorage. '
    "Se conecta mediante fetch() a un webhook Ngrok que deriva la consulta al motor N8N para procesamiento con IA.")

for item in [
    "Disponibilidad 24/7 sin intervención humana.",
    "Historial de conversación persistente entre páginas.",
    "Derivación automática al asesor si detecta intención de compra.",
]:
    agregar_bullet(doc, item)

doc.add_paragraph()

# Módulo 4
agregar_heading(doc, "3.4 Módulo 4 — Gestión y Consulta de Documentos", 2)
agregar_parrafo(doc,
    "Módulo de post-venta que permite consultar cotizaciones y ventas. Incluye un validador de RUT chileno "
    "en tiempo real basado en el algoritmo Módulo 11 del SII.")

agregar_heading(doc, "Algoritmo de Validación RUT (SII — Módulo 11):", 3)
for item in [
    "Serie de multiplicadores: [2, 3, 4, 5, 6, 7] aplicada al cuerpo del RUT en reversa.",
    "Cálculo del DV: 11 - (suma % 11).",
    "Casos especiales: resto 11 → DV '0' | resto 10 → DV 'K'.",
]:
    agregar_bullet(doc, item)

agregar_heading(doc, "Funcionalidades UX del validador:", 3)
for item in [
    "Auto-formateo en tiempo real: 12345678K → 12.345.678-K",
    "Feedback visual: borde verde/rojo + ícono check/cancel.",
    "Mensaje explicativo con el DV correcto si el RUT es inválido.",
    "Botón BUSCAR bloqueado hasta que el RUT sea válido.",
    "Modo Cotización: omite validación de RUT y acepta número de folio.",
]:
    agregar_bullet(doc, item)

doc.add_paragraph()

# Módulo 5 - Nav Móvil
agregar_heading(doc, "3.5 Módulo 5 — Navegación Responsiva (Menú Móvil)", 2)
agregar_parrafo(doc,
    "Menú hamburguesa implementado para pantallas menores a 1024px (lg). Al presionar el ícono ☰ "
    "se despliega un panel vertical con todos los enlaces del sitio incluyendo Buscar Documento, "
    "con íconos descriptivos y el botón destacado de Calculadora.")

doc.add_paragraph()

# ══════════════════════════════════════════
# SECCIÓN 4 – ARQUITECTURA TECNOLÓGICA
# ══════════════════════════════════════════
agregar_heading(doc, "4. Arquitectura Tecnológica", 1)
linea_separadora(doc)
doc.add_paragraph()

tabla_dos_col(doc,
    filas_data=[
        ("Presentación",       "HTML5 + CSS3 + Vanilla JS — Máxima velocidad sin frameworks pesados"),
        ("Estilos",            "TailwindCSS (CDN) — Diseño responsivo y modo oscuro"),
        ("Tipografía",         "Google Fonts — Inter, legibilidad premium"),
        ("Íconos",             "Material Symbols Outlined — Librería oficial Google"),
        ("Automatización",     "N8N (self-hosted o cloud) — Orquestador de flujos con soporte IA"),
        ("Proxy / Túnel",      "Ngrok — Exposición segura del servidor N8N vía HTTPS"),
        ("Scripts de mant.",   "Python 3 + Node.js — Actualización masiva de URLs y configuraciones"),
    ],
    header=["Capa", "Tecnología y Justificación"],
    col_widths=(4, 11)
)
doc.add_paragraph()

agregar_heading(doc, "Diagrama de Integración:", 2)
diagrama = doc.add_paragraph()
diagrama.paragraph_format.left_indent = Cm(1)
dr = diagrama.add_run(
    "[Navegador del Cliente]\n"
    "       │ HTTP/HTTPS\n"
    "       ▼\n"
    "[Sitio Web Estático — HTML/CSS/JS]\n"
    "       │ fetch() POST (JSON)\n"
    "       ▼\n"
    "[Ngrok Tunnel — URL HTTPS pública]\n"
    "       │ Forwarding\n"
    "       ▼\n"
    "[N8N Workflow — Servidor Local/Cloud]\n"
    "       │\n"
    "  ┌────┴────┐\n"
    "  ▼         ▼\n"
    "[Base de  [Notificación\n"
    " Datos /   WhatsApp /\n"
    " CRM]      Email]"
)
dr.font.name = 'Courier New'
dr.font.size = Pt(9)

doc.add_paragraph()

# ══════════════════════════════════════════
# SECCIÓN 5 – SCRIPTS DE MANTENIMIENTO
# ══════════════════════════════════════════
agregar_heading(doc, "5. Scripts de Mantenimiento y DevOps", 1)
linea_separadora(doc)
doc.add_paragraph()

tabla_dos_col(doc,
    filas_data=[
        ("update_ngrok.js / .py",   "Actualiza la URL de Ngrok en todos los archivos HTML del proyecto"),
        ("update_n8n.js / .py",     "Sincroniza los nodos y webhooks de N8N"),
        ("update_chat.py",          "Actualiza la configuración del chatbot masivamente"),
        ("add_seo.py",              "Inyecta meta-tags SEO en todas las páginas HTML"),
        ("inject_chatbot.py",       "Inserta el widget de chatbot en páginas que no lo tienen"),
        ("apply.js / apply2.js",    "Aplica cambios en lote sobre múltiples archivos"),
        ("generar_docx.py",         "Genera el presente documento de solución digital (.docx)"),
    ],
    header=["Script", "Función"],
    col_widths=(5, 10)
)
doc.add_paragraph()

# ══════════════════════════════════════════
# SECCIÓN 6 – FLUJOS CRÍTICOS
# ══════════════════════════════════════════
agregar_heading(doc, "6. Flujos Críticos de Usuario", 1)
linea_separadora(doc)
doc.add_paragraph()

flujos = [
    ("Flujo 1 — Navegación y Presupuesto", [
        "1. Cliente visita el sitio.",
        "2. Navega a la sección relevante (Jardines / Deportivo / Insumos).",
        "3. Usa la Calculadora para estimar su proyecto.",
        "4. Contacta via WhatsApp o formulario con datos claros.",
        "5. Asesor recibe un prospecto informado y cualificado.",
    ]),
    ("Flujo 2 — Consulta por Chatbot (Atención 24/7)", [
        "1. Cliente tiene una duda técnica o comercial.",
        "2. Abre el chatbot 'Queno' (disponible en todas las páginas).",
        "3. chatbot.js envía la consulta al webhook de Ngrok.",
        "4. N8N procesa la intención y genera la respuesta.",
        "5. Cliente recibe respuesta inmediata sin esperar a un asesor.",
        "6. [Opcional] N8N notifica al asesor si detecta intención de compra.",
    ]),
    ("Flujo 3 — Búsqueda de Documento (Posventa)", [
        "1. Cliente quiere revisar su cotización o venta.",
        "2. Accede a 'Buscar Documento' desde el menú de navegación.",
        "3. Selecciona el tipo: Cotización o Venta.",
        "4. Ingresa su RUT (validado en tiempo real con Módulo 11 SII) o número de folio.",
        "5. El sistema consulta la base de datos y retorna el documento.",
        "6. Cliente revisa el estado de su pedido de forma autónoma.",
    ]),
]

for titulo_flujo, pasos in flujos:
    agregar_heading(doc, titulo_flujo, 2)
    for paso in pasos:
        agregar_bullet(doc, paso)
    doc.add_paragraph()

# ══════════════════════════════════════════
# SECCIÓN 7 – MÉTRICAS DE ÉXITO
# ══════════════════════════════════════════
agregar_heading(doc, "7. Métricas de Éxito Esperadas", 1)
linea_separadora(doc)
doc.add_paragraph()

tabla_dos_col(doc,
    filas_data=[
        ("Tasa de Conversión",       "% visitantes que usan la calculadora y luego contactan   ≥ 15%"),
        ("Resolución en Chatbot",    "% de consultas resueltas sin escalar a asesor humano      ≥ 60%"),
        ("Disponibilidad del Sitio", "Tiempo de actividad del sitio y chatbot                   99.5%"),
        ("Tiempo de Respuesta",      "Tiempo promedio del chatbot para responder                 < 3 seg"),
        ("Leads Captados",           "Contactos nuevos generados mensualmente por el sistema     Medible vía N8N"),
    ],
    header=["Métrica", "Descripción y Objetivo"],
    col_widths=(4.5, 10.5)
)
doc.add_paragraph()

# ══════════════════════════════════════════
# PIE DE PÁGINA
# ══════════════════════════════════════════
linea_separadora(doc)
pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
pie_run = pie.add_run(
    f"© {datetime.date.today().year} Césped Sintético Chile — Documento generado automáticamente · "
    "Proyecto: CespedSintetico · gonzalo282/GameLevel"
)
pie_run.font.size = Pt(8)
pie_run.font.color.rgb = COLOR_GRIS
pie_run.font.italic = True

# ══════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════
doc.save(NOMBRE_ARCHIVO)
print(f"\n✅ Documento generado exitosamente: {NOMBRE_ARCHIVO}\n")
print(f"   Ubicación: {__import__('os').path.abspath(NOMBRE_ARCHIVO)}")
